# Consolidated Document Processors - All document processing logic

import os
import logging
from typing import Dict, Any, List, Optional, Tuple
import tempfile
import hashlib
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
import asyncio

from django.conf import settings
from django.core.files.storage import default_storage
from django.utils import timezone
from sentence_transformers import SentenceTransformer
import PyPDF2
import docx
from PIL import Image
import pytesseract
import groq

from .models import UserDocument, PublicDocument, ProcessingTask
from .database_consolidated import database_manager
from .security_consolidated import security_validator, encryption_manager

logger = logging.getLogger(__name__)

# ============================================================================
# BASE DOCUMENT PROCESSOR
# ============================================================================

class BaseDocumentProcessor:
    """Base class for document processing"""
    
    def __init__(self):
        self.embedding_model = None
        self.llm_client = None
        self.chunk_size = 1000
        self.chunk_overlap = 200
    
    def get_embedding_model(self):
        """Lazy load embedding model"""
        if self.embedding_model is None:
            self.embedding_model = SentenceTransformer('all-mpnet-base-v2')
        return self.embedding_model
    
    def get_llm_client(self):
        """Lazy load LLM client"""
        if self.llm_client is None:
            self.llm_client = groq.Groq(api_key=settings.GROQ_API_KEY)
        return self.llm_client
    
    def extract_text_from_file(self, file_path: str) -> Tuple[str, Optional[str]]:
        """Extract text from various file formats"""
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.txt':
                return self._extract_from_txt(file_path)
            elif file_ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff']:
                return self._extract_from_image(file_path)
            else:
                return "", f"Unsupported file type: {file_ext}"
                
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {str(e)}")
            return "", f"Text extraction failed: {str(e)}"
    
    def _extract_from_txt(self, file_path: str) -> Tuple[str, Optional[str]]:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return content, None
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    return content, None
                except UnicodeDecodeError:
                    continue
            return "", "Unable to decode text file"
    
    def _extract_from_pdf(self, file_path: str) -> Tuple[str, Optional[str]]:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
                        continue
            
            if not text.strip():
                return "", "No text found in PDF"
            
            return text, None
            
        except Exception as e:
            return "", f"PDF extraction failed: {str(e)}"
    
    def _extract_from_docx(self, file_path: str) -> Tuple[str, Optional[str]]:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        tables_text.append(" | ".join(row_text))
            
            # Combine all text
            all_text = "\n".join(paragraphs)
            if tables_text:
                all_text += "\n\n--- Tables ---\n" + "\n".join(tables_text)
            
            return all_text, None
            
        except Exception as e:
            return "", f"DOCX extraction failed: {str(e)}"
    
    def _extract_from_image(self, file_path: str) -> Tuple[str, Optional[str]]:
        """Extract text from image using OCR"""
        try:
            # Open and validate image
            image = Image.open(file_path)
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Extract text using Tesseract OCR
            extracted_text = pytesseract.image_to_string(image, lang='eng')
            
            if not extracted_text.strip():
                return "", "No text found in image"
            
            return extracted_text, None
            
        except Exception as e:
            return "", f"OCR extraction failed: {str(e)}"
    
    def chunk_text(self, text: str) -> List[str]:
        """Split text into chunks for processing"""
        if not text:
            return []
        
        # Simple sentence-based chunking
        sentences = text.split('. ')
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            # Add sentence to current chunk
            test_chunk = current_chunk + sentence + ". "
            
            if len(test_chunk) <= self.chunk_size:
                current_chunk = test_chunk
            else:
                # Save current chunk and start new one
                if current_chunk:
                    chunks.append(current_chunk.strip())
                
                # Handle overlap
                if len(chunks) > 0 and self.chunk_overlap > 0:
                    overlap_text = chunks[-1][-self.chunk_overlap:]
                    current_chunk = overlap_text + sentence + ". "
                else:
                    current_chunk = sentence + ". "
        
        # Add final chunk
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings for text chunks"""
        try:
            model = self.get_embedding_model()
            embeddings = model.encode(texts)
            return embeddings.tolist()
        except Exception as e:
            logger.error(f"Embedding generation failed: {str(e)}")
            return []
    
    def extract_entities(self, text: str) -> Dict[str, Any]:
        """Extract entities from text using LLM"""
        try:
            client = self.get_llm_client()
            
            prompt = f"""Extract key entities from this legal document text. Return a JSON object with the following structure:
{{
    "people": ["list of person names"],
    "organizations": ["list of organization names"],
    "locations": ["list of locations"],
    "dates": ["list of dates"],
    "legal_concepts": ["list of legal concepts"],
    "case_numbers": ["list of case numbers"],
    "statutes": ["list of statutes or laws mentioned"]
}}

Text: {text[:2000]}...

JSON:"""

            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
                max_tokens=1000
            )
            
            # Parse JSON response
            import json
            entities = json.loads(response.choices[0].message.content)
            return entities
            
        except Exception as e:
            logger.error(f"Entity extraction failed: {str(e)}")
            return {
                "people": [],
                "organizations": [],
                "locations": [],
                "dates": [],
                "legal_concepts": [],
                "case_numbers": [],
                "statutes": []
            }


# ============================================================================
# PUBLIC DOCUMENT PROCESSOR
# ============================================================================

class PublicDocumentProcessor(BaseDocumentProcessor):
    """Process public legal documents for the knowledge base"""
    
    def process_document(self, document: PublicDocument) -> Dict[str, Any]:
        """Process public document completely"""
        try:
            # Update status
            document.processing_status = 'processing'
            document.save()
            
            # Extract text
            text, error = self.extract_text_from_file(document.file_path.path)
            if error:
                document.processing_status = 'failed'
                document.save()
                return {'success': False, 'error': error}
            
            # Chunk text
            chunks = self.chunk_text(text)
            if not chunks:
                document.processing_status = 'failed'
                document.save()
                return {'success': False, 'error': 'No text chunks generated'}
            
            # Generate embeddings
            embeddings = self.generate_embeddings(chunks)
            if not embeddings:
                document.processing_status = 'failed'
                document.save()
                return {'success': False, 'error': 'Failed to generate embeddings'}
            
            # Extract entities
            entities = self.extract_entities(text)
            
            # Store in vector database
            vector_data = []
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                vector_data.append({
                    'document_id': str(document.id),
                    'chunk_id': f"{document.id}_{i}",
                    'content': chunk,
                    'embedding': embedding,
                    'metadata': {
                        'document_type': document.document_type,
                        'legal_domain': document.legal_domain,
                        'title': document.title,
                        'chunk_index': i,
                        'total_chunks': len(chunks)
                    }
                })
            
            # Insert into Milvus
            success = database_manager.milvus.insert_vectors(
                vector_data, 
                "public_documents", 
                is_personal=False
            )
            
            if not success:
                document.processing_status = 'failed'
                document.save()
                return {'success': False, 'error': 'Failed to store vectors'}
            
            # Store in graph database
            graph_data = {
                'id': str(document.id),
                'title': document.title,
                'content': text[:5000],  # First 5000 chars
                'document_type': document.document_type,
                'metadata': {
                    'legal_domain': document.legal_domain,
                    'jurisdiction': document.jurisdiction,
                    'entities': entities,
                    'chunk_count': len(chunks)
                }
            }
            
            graph_success = database_manager.neo4j.insert_document_graph(graph_data)
            
            # Update document status
            document.processing_status = 'completed'
            document.entities_extracted = entities
            document.embeddings_count = len(embeddings)
            document.processed_at = timezone.now()
            document.save()
            
            return {
                'success': True,
                'chunks_processed': len(chunks),
                'embeddings_generated': len(embeddings),
                'entities_extracted': len(sum(entities.values(), [])),
                'vector_storage': success,
                'graph_storage': graph_success
            }
            
        except Exception as e:
            logger.error(f"Public document processing failed: {str(e)}")
            document.processing_status = 'failed'
            document.save()
            return {'success': False, 'error': str(e)}


# ============================================================================
# PERSONAL DOCUMENT PROCESSOR
# ============================================================================

class PersonalDocumentProcessor(BaseDocumentProcessor):
    """Process personal user documents with data segregation"""
    
    def process_document(self, document: UserDocument) -> Dict[str, Any]:
        """Process personal document with user isolation"""
        try:
            # Update status
            document.status = 'processing'
            document.save()
            
            # Extract text
            text, error = self.extract_text_from_file(document.file_path.path)
            if error:
                document.status = 'failed'
                document.save()
                return {'success': False, 'error': error}
            
            # Encrypt sensitive content if required
            if getattr(settings, 'ENCRYPT_USER_DOCUMENTS', False):
                try:
                    encrypted_text = encryption_manager.encrypt_data(text)
                    # Store encrypted version separately
                    encrypted_path = document.file_path.path + '.encrypted'
                    with open(encrypted_path, 'w') as f:
                        f.write(encrypted_text)
                except Exception as e:
                    logger.warning(f"Document encryption failed: {str(e)}")
            
            # Chunk text
            chunks = self.chunk_text(text)
            if not chunks:
                document.status = 'failed'
                document.save()
                return {'success': False, 'error': 'No text chunks generated'}
            
            # Generate embeddings
            embeddings = self.generate_embeddings(chunks)
            if not embeddings:
                document.status = 'failed'
                document.save()
                return {'success': False, 'error': 'Failed to generate embeddings'}
            
            # Extract entities
            entities = self.extract_entities(text)
            
            # Store in user-specific vector collection
            user_collection_name = f"user_{document.user.id.hex}"
            vector_data = []
            
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                vector_data.append({
                    'user_id': document.user.id.hex,
                    'document_id': str(document.id),
                    'chunk_id': f"{document.id}_{i}",
                    'content': chunk,
                    'embedding': embedding,
                    'metadata': {
                        'file_name': document.file_name,
                        'file_type': document.file_type,
                        'chunk_index': i,
                        'total_chunks': len(chunks),
                        'upload_date': document.created_at.isoformat()
                    }
                })
            
            # Insert into user's Milvus collection
            success = database_manager.milvus.insert_vectors(
                vector_data, 
                user_collection_name, 
                is_personal=True
            )
            
            if not success:
                document.status = 'failed'
                document.save()
                return {'success': False, 'error': 'Failed to store vectors'}
            
            # Store in user's graph partition
            graph_data = {
                'id': str(document.id),
                'title': document.file_name,
                'content': text[:5000],  # First 5000 chars
                'metadata': {
                    'file_type': document.file_type,
                    'entities': entities,
                    'chunk_count': len(chunks)
                }
            }
            
            graph_success = database_manager.neo4j.insert_document_graph(
                graph_data, 
                user_id=document.user.id.hex
            )
            
            # Generate summary if requested
            summary = ""
            if document.summary_type:
                summary_result = self.generate_summary(text, document.summary_type)
                if summary_result['success']:
                    summary = summary_result['summary']
            
            # Update document status
            document.status = 'completed'
            document.summary = summary
            document.save()
            
            return {
                'success': True,
                'chunks_processed': len(chunks),
                'embeddings_generated': len(embeddings),
                'entities_extracted': len(sum(entities.values(), [])),
                'vector_storage': success,
                'graph_storage': graph_success,
                'summary_generated': bool(summary)
            }
            
        except Exception as e:
            logger.error(f"Personal document processing failed: {str(e)}")
            document.status = 'failed'
            document.save()
            return {'success': False, 'error': str(e)}
    
    def generate_summary(self, text: str, summary_type: str) -> Dict[str, Any]:
        """Generate document summary"""
        try:
            client = self.get_llm_client()
            
            # Truncate text if too long
            max_chars = 15000
            if len(text) > max_chars:
                text = text[:max_chars] + "\n\n[Document truncated due to length...]"
            
            # Summary prompts
            prompts = {
                'brief': f"Provide a brief 2-3 paragraph summary of this document:\n\n{text}",
                'comprehensive': f"Provide a comprehensive analysis including overview, key points, and implications:\n\n{text}",
                'legal_issues': f"Identify and analyze all legal issues in this document:\n\n{text}",
                'clause_by_clause': f"Provide a clause-by-clause analysis of this document:\n\n{text}"
            }
            
            prompt = prompts.get(summary_type, prompts['comprehensive'])
            
            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=2000
            )
            
            summary = response.choices[0].message.content
            
            return {'success': True, 'summary': summary}
            
        except Exception as e:
            logger.error(f"Summary generation failed: {str(e)}")
            return {'success': False, 'error': str(e)}


# ============================================================================
# BATCH PROCESSOR
# ============================================================================

class BatchDocumentProcessor:
    """Process multiple documents in parallel"""
    
    def __init__(self, max_workers: int = 3):
        self.max_workers = max_workers
        self.public_processor = PublicDocumentProcessor()
        self.personal_processor = PersonalDocumentProcessor()
    
    def process_public_documents(self, document_ids: List[str]) -> Dict[str, Any]:
        """Process multiple public documents in parallel"""
        results = {'successful': [], 'failed': []}
        
        try:
            documents = PublicDocument.objects.filter(id__in=document_ids)
            
            with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
                # Submit all tasks
                future_to_doc = {
                    executor.submit(self.public_processor.process_document, doc): doc
                    for doc in documents
                }
                
                # Collect results
                for future in as_completed(future_to_doc):
                    doc = future_to_doc[future]
                    try:
                        result = future.result()
                        if result['success']:
                            results['successful'].append({
                                'document_id': str(doc.id),
                                'title': doc.title,
                                'result': result
                            })
                        else:
                            results['failed'].append({
                                'document_id': str(doc.id),
                                'title': doc.title,
                                'error': result['error']
                            })
                    except Exception as e:
                        results['failed'].append({
                            'document_id': str(doc.id),
                            'title': doc.title,
                            'error': str(e)
                        })
            
            return results
            
        except Exception as e:
            logger.error(f"Batch processing failed: {str(e)}")
            return {'successful': [], 'failed': [], 'error': str(e)}
    
    def process_user_documents(self, user, document_ids: List[str]) -> Dict[str, Any]:
        """Process multiple user documents in parallel"""
        results = {'successful': [], 'failed': []}
        
        try:
            documents = UserDocument.objects.filter(
                id__in=document_ids, 
                user=user
            )
            
            with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
                # Submit all tasks
                future_to_doc = {
                    executor.submit(self.personal_processor.process_document, doc): doc
                    for doc in documents
                }
                
                # Collect results
                for future in as_completed(future_to_doc):
                    doc = future_to_doc[future]
                    try:
                        result = future.result()
                        if result['success']:
                            results['successful'].append({
                                'document_id': str(doc.id),
                                'file_name': doc.file_name,
                                'result': result
                            })
                        else:
                            results['failed'].append({
                                'document_id': str(doc.id),
                                'file_name': doc.file_name,
                                'error': result['error']
                            })
                    except Exception as e:
                        results['failed'].append({
                            'document_id': str(doc.id),
                            'file_name': doc.file_name,
                            'error': str(e)
                        })
            
            return results
            
        except Exception as e:
            logger.error(f"User batch processing failed: {str(e)}")
            return {'successful': [], 'failed': [], 'error': str(e)}


# ============================================================================
# CLEANUP PROCESSOR
# ============================================================================

class CleanupProcessor:
    """Handle cleanup of processed documents and data"""
    
    def cleanup_user_data(self, user_id: str) -> Dict[str, Any]:
        """Clean up all user data from databases"""
        try:
            results = {
                'milvus_cleanup': False,
                'neo4j_cleanup': False,
                'cache_cleanup': False
            }
            
            # Clean up Milvus data
            try:
                results['milvus_cleanup'] = database_manager.milvus.delete_user_data(user_id)
            except Exception as e:
                logger.error(f"Milvus cleanup failed: {str(e)}")
            
            # Clean up Neo4j data
            try:
                results['neo4j_cleanup'] = database_manager.neo4j.delete_user_graph_data(user_id)
            except Exception as e:
                logger.error(f"Neo4j cleanup failed: {str(e)}")
            
            # Clean up cache
            try:
                results['cache_cleanup'] = database_manager.cache.clear_user_cache(user_id)
            except Exception as e:
                logger.error(f"Cache cleanup failed: {str(e)}")
            
            return results
            
        except Exception as e:
            logger.error(f"User data cleanup failed: {str(e)}")
            return {'error': str(e)}
    
    def cleanup_failed_processing(self) -> Dict[str, Any]:
        """Clean up failed processing tasks"""
        try:
            # Find failed tasks older than 24 hours
            cutoff_time = timezone.now() - timezone.timedelta(hours=24)
            
            failed_tasks = ProcessingTask.objects.filter(
                status='failed',
                created_at__lt=cutoff_time
            )
            
            cleaned_count = 0
            for task in failed_tasks:
                try:
                    # Clean up associated data
                    if task.user_document:
                        task.user_document.status = 'failed'
                        task.user_document.save()
                    
                    if task.public_document:
                        task.public_document.processing_status = 'failed'
                        task.public_document.save()
                    
                    # Delete task
                    task.delete()
                    cleaned_count += 1
                    
                except Exception as e:
                    logger.warning(f"Failed to clean up task {task.id}: {str(e)}")
            
            return {
                'success': True,
                'cleaned_tasks': cleaned_count
            }
            
        except Exception as e:
            logger.error(f"Failed processing cleanup failed: {str(e)}")
            return {'success': False, 'error': str(e)}


# Create singleton instances
public_processor = PublicDocumentProcessor()
personal_processor = PersonalDocumentProcessor()
batch_processor = BatchDocumentProcessor()
cleanup_processor = CleanupProcessor()